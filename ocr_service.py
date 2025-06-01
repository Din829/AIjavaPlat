"""
AI Platform OCR Microservice

这个服务提供了PDF文档的OCR和结构化分析功能，使用多层处理策略：
1. 使用PyPDF2直接提取文本（适用于非扫描版PDF）
2. 使用Docling进行OCR和文档结构解析（适用于扫描版PDF）
3. 使用Google Gemini API进行高级文本理解和分析

服务提供RESTful API接口，支持文件上传和处理。
"""

# 启用 Docling 模型下载
import os
os.environ["DOCLING_ALLOW_DOWNLOADS"] = "1"
import sys
import uuid
import shutil
import json
import base64
import logging
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
from pathlib import Path
import fitz # PyMuPDF
from PIL import Image
import io

# FastAPI相关导入
from fastapi import FastAPI, UploadFile, File, HTTPException, Query, BackgroundTasks, Request, Form
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- 尝试导入必要的库 ---
# 1. PyPDF2 - 用于直接提取PDF文本
PYPDF2_AVAILABLE = False
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
    logger.info("Successfully imported PyPDF2.")
except ImportError:
    logger.warning("PyPDF2 not available. Direct PDF text extraction will be disabled.")

# 2. Docling - 用于OCR和文档结构解析
DOCLING_AVAILABLE = False
# 暂时禁用Docling以避免NumPy兼容性问题
# try:
#     from docling.document_converter import DocumentConverter
#     DOCLING_AVAILABLE = True
#     logger.info("Successfully imported DocumentConverter from docling.")
# except ImportError as e:
#     logger.warning(f"Failed to import Docling components: {e}")
#     logger.warning("Docling OCR processing will be disabled.")
logger.info("Docling temporarily disabled due to NumPy compatibility issues.")

# 3. Google Gemini API - 用于高级文本理解和分析
GEMINI_ENABLED = False
try:
    import google.generativeai as genai
    import PIL.Image  # 确保导入PIL.Image用于Gemini图像处理
    logger.info("Successfully imported google.generativeai.")
    GEMINI_ENABLED = True
except ImportError:
    logger.warning("google.generativeai module not found. Gemini enhancements will be disabled.")

# 4. Excel文件处理 - 用于处理Excel文档
EXCEL_PROCESSING_AVAILABLE = False
try:
    import pandas as pd
    import openpyxl
    EXCEL_PROCESSING_AVAILABLE = True
    logger.info("Successfully imported pandas and openpyxl for Excel processing.")
except ImportError as e:
    logger.warning(f"Failed to import Excel processing libraries: {e}")
    logger.warning("Excel file processing will be disabled.")

# 5. Word文档处理 - 用于处理Word文档（为未来扩展准备）
WORD_PROCESSING_AVAILABLE = False
try:
    import docx
    WORD_PROCESSING_AVAILABLE = True
    logger.info("Successfully imported python-docx for Word processing.")
except ImportError as e:
    logger.warning(f"Failed to import Word processing libraries: {e}")
    logger.warning("Word file processing will be disabled.")

# --- FastAPI 应用实例 ---
app = FastAPI(title="AI Platform OCR Microservice")

# --- Pydantic 模型定义 ---
# 简化的数据模型，专注于实用性和易用性

class ProcessingInfo(BaseModel):
    """处理信息，包括使用的处理方法和状态"""
    pypdf2_used: bool = False
    docling_used: bool = False
    gemini_used: bool = False
    force_ocr_used: bool = False
    processing_time_seconds: Optional[float] = None
    status: str = "success"
    error_message: Optional[str] = None

class DocumentMetadata(BaseModel):
    """文档元数据"""
    original_filename: Optional[str] = None
    processed_at: Optional[str] = None
    page_count: Optional[int] = None
    source_format: Optional[str] = None
    language: Optional[str] = None
    title: Optional[str] = None
    author: Optional[str] = None
    creation_date: Optional[str] = None

class TableInfo(BaseModel):
    """表格信息"""
    table_id: str
    page_number: int
    title: Optional[str] = None
    headers: Optional[List[str]] = None
    rows: List[List[str]] = []
    raw_text: Optional[Union[str, Dict[str, Any], List[Dict[str, Any]]]] = None

class ImageInfo(BaseModel):
    """图像信息"""
    image_id: str
    page_number: int
    description: Optional[str] = None
    ocr_text: Optional[str] = None
    mime_type: Optional[str] = None
    data: Optional[str] = None  # Base64编码的图像数据

class PageContent(BaseModel):
    """页面内容"""
    page_number: int
    text: str = ""
    tables: List[TableInfo] = []
    images: List[ImageInfo] = []

class GeminiAnalysis(BaseModel):
    """Gemini分析结果"""
    summary: Optional[str] = None
    key_points: List[str] = []
    structured_data: Optional[Dict[str, Any]] = None
    translation: Optional[str] = None
    raw_response: Optional[str] = None

class OcrResponse(BaseModel):
    """OCR响应"""
    document_metadata: DocumentMetadata
    processing_info: ProcessingInfo
    pages: List[PageContent] = []
    full_text: str = ""
    tables: List[TableInfo] = []
    images: List[ImageInfo] = []
    gemini_analysis: Optional[GeminiAnalysis] = None


# --- 全局变量与常量 ---
# 上传文件临时目录
UPLOAD_DIR = "temp_uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 创建本地缓存目录
LOCAL_CACHE_DIR = "docling_cache"
os.makedirs(LOCAL_CACHE_DIR, exist_ok=True)

# 设置环境变量，告诉Hugging Face和Docling使用本地缓存目录
os.environ["HF_HOME"] = os.path.abspath(LOCAL_CACHE_DIR)
os.environ["TRANSFORMERS_CACHE"] = os.path.join(os.path.abspath(LOCAL_CACHE_DIR), "transformers")
os.environ["HF_DATASETS_CACHE"] = os.path.join(os.path.abspath(LOCAL_CACHE_DIR), "datasets")
os.environ["HUGGINGFACE_HUB_CACHE"] = os.path.join(os.path.abspath(LOCAL_CACHE_DIR), "hub")
os.environ["DOCLING_CACHE_DIR"] = os.path.abspath(LOCAL_CACHE_DIR)

# 创建所有必要的子目录
for subdir in ["transformers", "datasets", "hub"]:
    os.makedirs(os.path.join(LOCAL_CACHE_DIR, subdir), exist_ok=True)

logger.info(f"Using local cache directory: {os.path.abspath(LOCAL_CACHE_DIR)}")

# --- Gemini API配置 ---
# 从环境变量获取Gemini API密钥
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyDFLyEYqgaC6plSFF5IjvQEW0FEug6o14o")

# 定义可用的Gemini模型配置
GEMINI_MODELS = {
    "gemini-2.5-pro-preview-05-06": {
        "name": "models/gemini-2.5-pro-preview-05-06",
        "display_name": "Gemini 2.5 Pro Preview 05-06 (最佳OCR质量)",
        "description": "专门优化用于OCR和图像文字识别的模型，精度最高但处理时间较长",
        "speed": "slow",
        "quality": "highest"
    },
    "gemini-2.5-flash-preview-05-20": {
        "name": "models/gemini-2.5-flash-preview-05-20",
        "display_name": "Gemini 2.5 Flash Preview 05-20 (快速)",
        "description": "最新的快速模型，在保持良好质量的同时大幅提升处理速度",
        "speed": "fast",
        "quality": "good"
    }
}

# 设置默认模型（优先使用Flash模型以提升速度）
DEFAULT_GEMINI_MODEL_KEY = "gemini-2.5-flash-preview-05-20"
DEFAULT_GEMINI_MODEL = None
AVAILABLE_GEMINI_MODELS = []

# 配置Gemini API
if GEMINI_ENABLED:
    try:
        # 尝试使用最新的API版本
        genai.configure(api_key=GEMINI_API_KEY, transport="rest")
        logger.info("Configured Gemini API with REST transport")

        # 检查可用的模型
        try:
            models = genai.list_models()
            model_names = [m.name for m in models]
            logger.info(f"Available Gemini models: {model_names}")

            # 检查哪些预定义模型可用
            for key, config in GEMINI_MODELS.items():
                if config["name"] in model_names:
                    AVAILABLE_GEMINI_MODELS.append(key)
                    logger.info(f"Model available: {key} -> {config['name']}")

            # 设置默认模型
            if DEFAULT_GEMINI_MODEL_KEY in AVAILABLE_GEMINI_MODELS:
                DEFAULT_GEMINI_MODEL = GEMINI_MODELS[DEFAULT_GEMINI_MODEL_KEY]["name"]
            elif AVAILABLE_GEMINI_MODELS:
                # 如果默认模型不可用，使用第一个可用模型
                DEFAULT_GEMINI_MODEL_KEY = AVAILABLE_GEMINI_MODELS[0]
                DEFAULT_GEMINI_MODEL = GEMINI_MODELS[DEFAULT_GEMINI_MODEL_KEY]["name"]
            else:
                # 如果没有预定义模型可用，使用备选模型
                for model_name in ["models/gemini-1.5-pro", "models/gemini-1.5-flash", "models/gemini-pro"]:
                    if model_name in model_names:
                        DEFAULT_GEMINI_MODEL = model_name
                        break

            logger.info(f"Using default Gemini model: {DEFAULT_GEMINI_MODEL}")
            logger.info(f"Available model keys: {AVAILABLE_GEMINI_MODELS}")
        except Exception as e:
            logger.warning(f"Failed to list available Gemini models: {e}")
            DEFAULT_GEMINI_MODEL = "models/gemini-pro"
    except Exception as e:
        logger.error(f"Failed to configure Gemini API: {e}")
        GEMINI_ENABLED = False
else:
    DEFAULT_GEMINI_MODEL = None

# --- 默认提示词 ---
# 完整版提示词（用于高质量模型）
DEFAULT_PDF_PROMPT = """
请分析这个PDF文档，并提取以下信息：
1. 文档的标题和主题
2. 文档的主要内容和结构
3. 文档中的关键信息和数据
4. 表格内容（如有）
5. 图表或图像描述（如有）

请以JSON格式返回结果：
{
    "title": "文档标题",
    "document_type": "文档类型",
    "language": "文档语言",
    "summary": "文档摘要",
    "key_points": ["关键点1", "关键点2", ...],
    "tables": [{"table_title": "表格标题", "content": "表格内容"}],
    "images": [{"description": "图像描述"}],
    "translation": "文档内容的中文翻译（如果原文不是中文）"
}
"""

# 简化版提示词（用于快速模型，专注核心任务）
SIMPLIFIED_PDF_PROMPT = """
请分析以下文档内容并提供简要总结：

请以JSON格式返回：
{
    "title": "文档标题",
    "summary": "文档主要内容摘要",
    "key_points": ["关键点1", "关键点2", "关键点3"],
    "language": "文档语言"
}
"""

# 日语文档的默认提示词
DEFAULT_JAPANESE_PDF_PROMPT = """
你是一个专业的日语文档分析专家。请分析这个日语PDF文档，并提取以下信息：

1. 文档的标题和主题
2. 文档的主要内容和结构
3. 文档中的关键信息和数据
4. 表格内容（如有）
5. 图表或图像描述（如有）

请以JSON格式返回结果：
{
    "title": "文档标题",
    "document_type": "文档类型",
    "language": "文档语言",
    "summary": "文档摘要",
    "key_fields": [{"field_name": "名称", "field_value": "值"}],
    "tables": [{"table_title": "表格标题", "content": "表格内容"}],
    "images": [{"description": "图像描述"}],
    "translation": "文档内容的中文翻译"
}
"""

# 日语文档的简化提示词
SIMPLIFIED_JAPANESE_PDF_PROMPT = """
请分析以下日语文档内容并提供简要总结：

请以JSON格式返回：
{
    "title": "文档标题",
    "summary": "文档主要内容摘要",
    "key_fields": [{"field_name": "名称", "field_value": "值"}],
    "translation": "文档内容的中文翻译"
}
"""

# --- 辅助函数 ---

def cleanup_temp_file(file_path: str):
    """删除临时文件"""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        logger.error(f"Error cleaning up temp file {file_path}: {e}")

def extract_images_from_pdf(file_path: str) -> List[ImageInfo]:
    """
    使用PyMuPDF从PDF文件中提取嵌入的图像

    Args:
        file_path: PDF文件路径

    Returns:
        包含图像信息的列表
    """
    images = []

    try:
        # 使用PyMuPDF打开PDF文件
        doc = fitz.open(file_path)
        logger.info(f"开始从PDF提取图像，共{len(doc)}页")

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            # 获取页面中的图像列表
            image_list = page.get_images()
            logger.info(f"第{page_num + 1}页发现{len(image_list)}个图像")

            for img_index, img in enumerate(image_list):
                try:
                    # 获取图像引用
                    xref = img[0]

                    # 提取图像数据
                    pix = fitz.Pixmap(doc, xref)

                    # 检查图像格式，避免CMYK格式问题
                    if pix.n - pix.alpha < 4:  # 确保不是CMYK格式
                        # 转换为PNG格式的字节数据
                        if pix.alpha:
                            # 如果有透明通道，保持PNG格式
                            img_data = pix.tobytes("png")
                            mime_type = "image/png"
                        else:
                            # 如果没有透明通道，可以使用JPEG格式（更小）
                            img_data = pix.tobytes("jpeg")
                            mime_type = "image/jpeg"

                        # 转换为Base64编码
                        img_base64 = base64.b64encode(img_data).decode('utf-8')

                        # 创建图像信息对象
                        image_info = ImageInfo(
                            image_id=f"page_{page_num + 1}_img_{img_index + 1}",
                            page_number=page_num + 1,
                            description=f"第{page_num + 1}页的图像{img_index + 1}",
                            mime_type=mime_type,
                            data=img_base64
                        )

                        images.append(image_info)
                        logger.info(f"成功提取图像: {image_info.image_id}, 大小: {len(img_data)} bytes")
                    else:
                        logger.warning(f"跳过CMYK格式图像: 第{page_num + 1}页图像{img_index + 1}")

                    # 释放Pixmap内存
                    pix = None

                except Exception as e:
                    logger.error(f"提取第{page_num + 1}页图像{img_index + 1}时出错: {e}")
                    continue

        doc.close()
        logger.info(f"图像提取完成，共提取{len(images)}个图像")

    except Exception as e:
        logger.error(f"从PDF提取图像时发生错误: {e}")
        return []

    return images

def extract_text_with_pypdf2(file_path: str) -> Dict[str, Any]:
    """
    使用PyPDF2直接从PDF文件中提取文本

    Args:
        file_path: PDF文件路径

    Returns:
        包含提取文本的字典
    """
    if not PYPDF2_AVAILABLE:
        return {"error": "PyPDF2 not available"}

    try:
        # 打开PDF文件
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)

            # 获取页面数量
            num_pages = len(pdf_reader.pages)
            logger.info(f"PDF has {num_pages} pages")

            # 提取文本
            all_text = []
            pages = []

            for i in range(num_pages):
                page = pdf_reader.pages[i]
                text = page.extract_text() or ""

                # 创建页面数据
                page_data = {
                    "page_number": i+1,
                    "text": text
                }

                pages.append(page_data)
                all_text.append(text)

            # 创建结果
            result = {
                "document_type": "PDF",
                "processed_at": datetime.now().isoformat(),
                "original_filename": os.path.basename(file_path),
                "page_count": num_pages,
                "pages": pages,
                "full_text": "\n".join(all_text)
            }

            # 尝试提取文档信息
            if hasattr(pdf_reader, 'metadata') and pdf_reader.metadata:
                metadata = pdf_reader.metadata
                result["title"] = metadata.get('/Title', None)
                result["author"] = metadata.get('/Author', None)
                result["creation_date"] = metadata.get('/CreationDate', None)

            return result

    except Exception as e:
        logger.exception(f"Error extracting text with PyPDF2: {e}")
        return {"error": f"Error extracting text with PyPDF2: {str(e)}"}

def process_with_docling(file_path: str, force_ocr: bool = False) -> Dict[str, Any]:
    """
    使用Docling处理PDF文件

    Args:
        file_path: 要处理的PDF文件路径
        force_ocr: 是否强制使用OCR，即使PDF包含文本层

    Returns:
        包含提取文本和元数据的字典
    """
    import os # 确保 os 在此函数作用域内可用，以解决 linter 可能的误报
    if not DOCLING_AVAILABLE:
        return {"error": "Docling not available"}

    try:
        # 尝试导入并配置pipeline_options
        try:
            # 根据 Docling 2.32.0 版本的文档，正确的导入路径是 docling.datamodel.pipeline_options
            try:
                from docling.datamodel.pipeline_options import PdfPipelineOptions, EasyOcrOptions
                logger.info("Successfully imported from docling.datamodel.pipeline_options")

                # 使用 PdfPipelineOptions 更针对PDF处理
                pipeline_options = PdfPipelineOptions()

                # 设置缓存目录 - 在 Docling 2.32.0 中，可能使用不同的属性名
                # 尝试设置缓存目录，如果属性不存在则忽略
                try:
                    pipeline_options.cache_dir = os.path.abspath(LOCAL_CACHE_DIR)
                except Exception as e:
                    logger.info(f"Could not set cache_dir: {e}. This is normal for newer Docling versions.")
                    # 尝试其他可能的属性名
                    try:
                        pipeline_options.artifacts_path = os.path.abspath(LOCAL_CACHE_DIR)
                    except Exception as e2:
                        logger.info(f"Could not set artifacts_path: {e2}. Using default cache location.")

                # 启用OCR，并根据参数决定是否强制使用OCR
                # 尝试设置OCR选项，如果属性不存在则忽略
                try:
                    pipeline_options.do_ocr = True
                except Exception as e:
                    logger.info(f"Could not set do_ocr: {e}. This is normal for newer Docling versions.")

                try:
                    pipeline_options.force_ocr = force_ocr
                except Exception as e:
                    logger.info(f"Could not set force_ocr: {e}. This is normal for newer Docling versions.")

                # 尝试配置OCR选项，特别是对日语的支持
                ocr_options = EasyOcrOptions()
                # 添加日语支持 - 在不同版本中，属性名可能不同
                try:
                    # 尝试设置 languages 属性
                    ocr_options.languages = ["ja", "en"]
                except Exception as e:
                    logger.info(f"Could not set languages: {e}. Trying alternative attribute names.")
                    # 尝试设置 lang 属性
                    try:
                        ocr_options.lang = ["ja", "en"]
                    except Exception as e2:
                        logger.info(f"Could not set lang: {e2}. Trying alternative attribute names.")
                        # 尝试设置 language 属性
                        try:
                            ocr_options.language = ["ja", "en"]
                        except Exception as e3:
                            logger.info(f"Could not set language: {e3}. Using default language settings.")

                # 尝试设置 OCR 选项
                try:
                    pipeline_options.ocr_options = ocr_options
                except Exception as e:
                    logger.info(f"Could not set ocr_options: {e}. Using default OCR settings.")
                logger.info("Configured EasyOCR with Japanese and English language support")
            except ImportError:
                # 尝试其他可能的导入路径
                logger.info("Trying alternative import paths for Docling 2.32.0")
                try:
                    # 尝试从 docling.pipeline_options 导入
                    from docling.pipeline_options import PdfPipelineOptions, EasyOcrOptions
                    logger.info("Successfully imported from docling.pipeline_options")

                    pipeline_options = PdfPipelineOptions()
                    # 设置缓存目录 - 在 Docling 2.32.0 中，可能使用不同的属性名
                    # 尝试设置缓存目录，如果属性不存在则忽略
                    try:
                        pipeline_options.cache_dir = os.path.abspath(LOCAL_CACHE_DIR)
                    except Exception as e:
                        logger.info(f"Could not set cache_dir: {e}. This is normal for newer Docling versions.")
                        # 尝试其他可能的属性名
                        try:
                            pipeline_options.artifacts_path = os.path.abspath(LOCAL_CACHE_DIR)
                        except Exception as e2:
                            logger.info(f"Could not set artifacts_path: {e2}. Using default cache location.")
                    # 尝试设置OCR选项，如果属性不存在则忽略
                    try:
                        pipeline_options.do_ocr = True
                    except Exception as e:
                        logger.info(f"Could not set do_ocr: {e}. This is normal for newer Docling versions.")

                    try:
                        pipeline_options.force_ocr = force_ocr
                    except Exception as e:
                        logger.info(f"Could not set force_ocr: {e}. This is normal for newer Docling versions.")

                    # 尝试配置OCR选项，特别是对日语的支持
                    ocr_options = EasyOcrOptions()
                    # 添加日语支持 - 在不同版本中，属性名可能不同
                    try:
                        # 尝试设置 languages 属性
                        ocr_options.languages = ["ja", "en"]
                    except Exception as e:
                        logger.info(f"Could not set languages: {e}. Trying alternative attribute names.")
                        # 尝试设置 lang 属性
                        try:
                            ocr_options.lang = ["ja", "en"]
                        except Exception as e2:
                            logger.info(f"Could not set lang: {e2}. Trying alternative attribute names.")
                            # 尝试设置 language 属性
                            try:
                                ocr_options.language = ["ja", "en"]
                            except Exception as e3:
                                logger.info(f"Could not set language: {e3}. Using default language settings.")

                    # 尝试设置 OCR 选项
                    try:
                        pipeline_options.ocr_options = ocr_options
                    except Exception as e:
                        logger.info(f"Could not set ocr_options: {e}. Using default OCR settings.")
                    logger.info("Configured EasyOCR with Japanese and English language support (alternative path)")
                except ImportError:
                    # 尝试从 docling.core.pipeline_options 导入
                    try:
                        from docling.core.pipeline_options import PdfPipelineOptions, EasyOcrOptions
                        logger.info("Successfully imported from docling.core.pipeline_options")

                        pipeline_options = PdfPipelineOptions()
                        # 设置缓存目录 - 在 Docling 2.32.0 中，可能使用不同的属性名
                        # 尝试设置缓存目录，如果属性不存在则忽略
                        try:
                            pipeline_options.cache_dir = os.path.abspath(LOCAL_CACHE_DIR)
                        except Exception as e:
                            logger.info(f"Could not set cache_dir: {e}. This is normal for newer Docling versions.")
                            # 尝试其他可能的属性名
                            try:
                                pipeline_options.artifacts_path = os.path.abspath(LOCAL_CACHE_DIR)
                            except Exception as e2:
                                logger.info(f"Could not set artifacts_path: {e2}. Using default cache location.")
                        # 尝试设置OCR选项，如果属性不存在则忽略
                        try:
                            pipeline_options.do_ocr = True
                        except Exception as e:
                            logger.info(f"Could not set do_ocr: {e}. This is normal for newer Docling versions.")

                        try:
                            pipeline_options.force_ocr = force_ocr
                        except Exception as e:
                            logger.info(f"Could not set force_ocr: {e}. This is normal for newer Docling versions.")

                        # 尝试配置OCR选项，特别是对日语的支持
                        ocr_options = EasyOcrOptions()
                        # 添加日语支持 - 在不同版本中，属性名可能不同
                        try:
                            # 尝试设置 languages 属性
                            ocr_options.languages = ["ja", "en"]
                        except Exception as e:
                            logger.info(f"Could not set languages: {e}. Trying alternative attribute names.")
                            # 尝试设置 lang 属性
                            try:
                                ocr_options.lang = ["ja", "en"]
                            except Exception as e2:
                                logger.info(f"Could not set lang: {e2}. Trying alternative attribute names.")
                                # 尝试设置 language 属性
                                try:
                                    ocr_options.language = ["ja", "en"]
                                except Exception as e3:
                                    logger.info(f"Could not set language: {e3}. Using default language settings.")

                        # 尝试设置 OCR 选项
                        try:
                            pipeline_options.ocr_options = ocr_options
                        except Exception as e:
                            logger.info(f"Could not set ocr_options: {e}. Using default OCR settings.")
                        logger.info("Configured EasyOCR with Japanese and English language support (alternative path)")
                    except ImportError:
                        # 尝试从 docling_core.pipeline_options 导入
                        from docling_core.pipeline_options import PdfPipelineOptions, EasyOcrOptions
                        logger.info("Successfully imported from docling_core.pipeline_options")

                        pipeline_options = PdfPipelineOptions()
                        # 设置缓存目录 - 在 Docling 2.32.0 中，可能使用不同的属性名
                        # 尝试设置缓存目录，如果属性不存在则忽略
                        try:
                            pipeline_options.cache_dir = os.path.abspath(LOCAL_CACHE_DIR)
                        except Exception as e:
                            logger.info(f"Could not set cache_dir: {e}. This is normal for newer Docling versions.")
                            # 尝试其他可能的属性名
                            try:
                                pipeline_options.artifacts_path = os.path.abspath(LOCAL_CACHE_DIR)
                            except Exception as e2:
                                logger.info(f"Could not set artifacts_path: {e2}. Using default cache location.")
                        # 尝试设置OCR选项，如果属性不存在则忽略
                        try:
                            pipeline_options.do_ocr = True
                        except Exception as e:
                            logger.info(f"Could not set do_ocr: {e}. This is normal for newer Docling versions.")

                        try:
                            pipeline_options.force_ocr = force_ocr
                        except Exception as e:
                            logger.info(f"Could not set force_ocr: {e}. This is normal for newer Docling versions.")

                        # 尝试配置OCR选项，特别是对日语的支持
                        ocr_options = EasyOcrOptions()
                        # 添加日语支持 - 在不同版本中，属性名可能不同
                        try:
                            # 尝试设置 languages 属性
                            ocr_options.languages = ["ja", "en"]
                        except Exception as e:
                            logger.info(f"Could not set languages: {e}. Trying alternative attribute names.")
                            # 尝试设置 lang 属性
                            try:
                                ocr_options.lang = ["ja", "en"]
                            except Exception as e2:
                                logger.info(f"Could not set lang: {e2}. Trying alternative attribute names.")
                                # 尝试设置 language 属性
                                try:
                                    ocr_options.language = ["ja", "en"]
                                except Exception as e3:
                                    logger.info(f"Could not set language: {e3}. Using default language settings.")

                        # 尝试设置 OCR 选项
                        try:
                            pipeline_options.ocr_options = ocr_options
                        except Exception as e:
                            logger.info(f"Could not set ocr_options: {e}. Using default OCR settings.")
                        logger.info("Configured EasyOCR with Japanese and English language support (alternative path)")

            logger.info(f"Using pipeline options: {type(pipeline_options)}, OCR enabled: {getattr(pipeline_options, 'do_ocr', 'N/A')}, force OCR: {getattr(pipeline_options, 'force_ocr', 'N/A')}")

            # 在 Docling 2.32.0 中，需要使用 format_options 参数
            from docling.datamodel.base_models import InputFormat
            from docling.document_converter import PdfFormatOption

            # 启用模型下载
            import os
            os.environ["DOCLING_ALLOW_DOWNLOADS"] = "1"

            # 创建 DocumentConverter 实例
            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                }
            )
        except ImportError as e_import: # 捕获导入错误
            logger.warning(f"Failed to import Docling pipeline options (ImportError: {e_import}). Using default DocumentConverter options. THIS IS LIKELY DUE TO AN INSTALLATION ISSUE.")
            # 使用默认的 DocumentConverter，不指定任何特殊选项
            converter = DocumentConverter()
        except Exception as e_config: # 捕获其他配置错误
            logger.warning(f"Failed to configure custom Docling pipeline options (Exception: {e_config}). Using default DocumentConverter options.")
            # 使用默认的 DocumentConverter，不指定任何特殊选项
            converter = DocumentConverter()

        # 开始处理PDF
        logger.info(f"Starting Docling conversion for '{file_path}'")
        start_time = datetime.now()
        docling_result = converter.convert(source=file_path)
        end_time = datetime.now()
        conversion_time = (end_time - start_time).total_seconds()
        logger.info(f"Docling conversion finished in {conversion_time:.2f} seconds")

        # 检查转换结果
        if not docling_result:
            logger.error(f"Docling processing failed: No result returned")
            return {"error": "Docling processing failed: No result returned"}

        # 检查转换状态
        if hasattr(docling_result, 'status') and hasattr(docling_result.status, 'name'):
            if docling_result.status.name != "SUCCESS":
                status_msg = getattr(docling_result, 'status_message', "Unknown error")
                logger.error(f"Docling processing failed: {status_msg}")
                return {"error": f"Docling processing error: {status_msg}"}

        # 获取文档对象
        document = getattr(docling_result, 'document', None)
        if not document:
            logger.error("Docling conversion successful but no document object returned")
            return {"error": "Docling error: No document object in result"}

        # 创建结果字典
        result = {
            "document_type": "PDF",
            "processed_at": datetime.now().isoformat(),
            "original_filename": os.path.basename(file_path),
            "conversion_time": conversion_time
        }

        # 使用正确的Docling API提取文本
        # 首先尝试使用export_to_markdown获取完整文本
        try:
            full_text = document.export_to_markdown()
            result["full_text"] = full_text
            logger.info(f"Successfully extracted text using export_to_markdown, length: {len(full_text)}")
        except Exception as e:
            logger.warning(f"Failed to use export_to_markdown: {e}")
            result["full_text"] = ""

        # 如果export_to_markdown没有提取到文本，尝试其他方法
        if not result["full_text"].strip():
            try:
                # 尝试使用export_to_text方法
                if hasattr(document, 'export_to_text'):
                    full_text = document.export_to_text()
                    result["full_text"] = full_text
                    logger.info(f"Successfully extracted text using export_to_text, length: {len(full_text)}")
            except Exception as e:
                logger.warning(f"Failed to use export_to_text: {e}")

        # 如果仍然没有文本，尝试直接访问文档内容
        if not result["full_text"].strip():
            try:
                # 尝试访问document.text属性
                if hasattr(document, 'text'):
                    full_text = document.text
                    result["full_text"] = full_text
                    logger.info(f"Successfully extracted text using document.text, length: {len(full_text)}")
            except Exception as e:
                logger.warning(f"Failed to use document.text: {e}")

        # 尝试获取页面信息
        result_pages = []
        all_text = []

        # 检查文档是否有页面 - 修复页面访问问题
        page_count = 0
        if hasattr(document, 'pages'):
            try:
                # 尝试获取页面数量
                if isinstance(document.pages, (list, tuple)):
                    page_count = len(document.pages)
                elif hasattr(document.pages, '__len__'):
                    page_count = len(document.pages)
                else:
                    # 如果pages不是列表，可能是页面数量
                    page_count = int(document.pages) if isinstance(document.pages, (int, float)) else 0

                result["page_count"] = page_count
                logger.info(f"Document has {page_count} pages")
            except Exception as e:
                logger.warning(f"Failed to get page count: {e}")
                page_count = 1
                result["page_count"] = 1

        # 如果没有提取到文本且有页面信息，尝试从页面元素中提取
        if not result["full_text"].strip() and page_count > 0:
            try:
                # 尝试从文档的内部结构中提取文本
                if hasattr(document, '_elements') and document._elements:
                    elements_text = []
                    for element in document._elements:
                        if hasattr(element, 'text') and element.text:
                            elements_text.append(element.text)
                    if elements_text:
                        result["full_text"] = "\n".join(elements_text)
                        logger.info(f"Successfully extracted text from document elements, length: {len(result['full_text'])}")

                # 尝试从文档的其他属性中提取文本
                elif hasattr(document, 'content') and document.content:
                    result["full_text"] = str(document.content)
                    logger.info(f"Successfully extracted text from document.content, length: {len(result['full_text'])}")

            except Exception as e:
                logger.warning(f"Failed to extract text from document elements: {e}")

        # 创建页面结构
        if page_count > 0:
            # 如果有完整文本，尝试按页面分割
            if result["full_text"].strip():
                text_per_page = len(result["full_text"]) // page_count if page_count > 0 else len(result["full_text"])
                for page_num in range(page_count):
                    start_idx = page_num * text_per_page
                    end_idx = (page_num + 1) * text_per_page if page_num < page_count - 1 else len(result["full_text"])
                    page_text = result["full_text"][start_idx:end_idx] if start_idx < len(result["full_text"]) else ""

                    result_pages.append({
                        "page_number": page_num + 1,
                        "text": page_text,
                        "elements": []
                    })
            else:
                # 没有文本，创建空页面
                for page_num in range(page_count):
                    result_pages.append({
                        "page_number": page_num + 1,
                        "text": "",
                        "elements": []
                    })
        else:
            # 没有页面信息，创建单页结构
            result["page_count"] = 1
            result_pages.append({
                "page_number": 1,
                "text": result["full_text"],
                "elements": []
            })

        result["pages"] = result_pages

        # ---- START DEBUG PRINTS ----
        logger.info(f"Docling Result (inside process_with_docling): {result}")
        if hasattr(document, 'pages') and document.pages:
            for i, page_obj in enumerate(document.pages):
                logger.info(f"  Page {i+1} Object Type: {type(page_obj)}")
                logger.info(f"  Page {i+1} Object Attributes: {dir(page_obj)}")
                if hasattr(page_obj, 'text_content'): # Common attribute for page text
                    logger.info(f"  Page {i+1} Text (from .text_content): {str(page_obj.text_content)[:500]}")
                elif hasattr(page_obj, 'text'):
                    logger.info(f"  Page {i+1} Text (from .text): {str(page_obj.text)[:500]}")

                # Log elements within the page to see their text attributes
                if hasattr(page_obj, 'elements') and page_obj.elements:
                    logger.info(f"  Page {i+1} has {len(page_obj.elements)} elements.")
                    for elem_idx, elem in enumerate(page_obj.elements):
                        elem_text_parts_debug = []
                        if hasattr(elem, 'text'):
                            elem_text_parts_debug.append(f"elem.text: {str(elem.text)}")
                        if hasattr(elem, 'texts'):
                             elem_text_parts_debug.append(f"elem.texts: {[str(t.text) for t in elem.texts if hasattr(t, 'text')]}")
                        if elem_text_parts_debug:
                            logger.info(f"    Element {elem_idx + 1} Text Debug: {' | '.join(elem_text_parts_debug)[:200]}")
        # ---- END DEBUG PRINTS ----

        return result

    except Exception as e:
        logger.exception(f"Error processing file with Docling: {e}")
        return {"error": f"Error processing file with Docling: {str(e)}"}

# --- 新增优化的Gemini处理函数 ---

def get_prompt_for_model(model_key: str, language: str = "auto") -> str:
    """
    根据模型类型和语言选择合适的提示词

    Args:
        model_key: 模型键值 (gemini-2.5-pro, gemini-1.5-pro, gemini-1.5-flash)
        language: 文档语言

    Returns:
        合适的提示词
    """
    # 判断是否使用简化提示词（Flash模型使用简化版本）
    use_simplified = model_key == "gemini-1.5-flash"

    if language.lower() == "ja" or language.lower() == "japanese":
        return SIMPLIFIED_JAPANESE_PDF_PROMPT if use_simplified else DEFAULT_JAPANESE_PDF_PROMPT
    else:
        return SIMPLIFIED_PDF_PROMPT if use_simplified else DEFAULT_PDF_PROMPT

def truncate_text(text: str, max_length: int = 10000) -> str:
    """
    截断文本到指定长度，避免超长输入

    Args:
        text: 原始文本
        max_length: 最大长度

    Returns:
        截断后的文本
    """
    if len(text) <= max_length:
        return text

    # 截断并添加省略号
    truncated = text[:max_length]
    # 尝试在单词边界截断
    last_space = truncated.rfind(' ')
    if last_space > max_length * 0.8:  # 如果最后一个空格位置合理
        truncated = truncated[:last_space]

    return truncated + "...\n[文本已截断]"

def analyze_text_with_gemini_sync(text: str, model_key: str = None, language: str = "auto") -> Dict[str, Any]:
    """
    使用Gemini分析文本内容（同步版本，用于非异步环境）

    Args:
        text: 要分析的文本内容
        model_key: 使用的模型键值，如果为None则使用默认模型
        language: 文档语言

    Returns:
        包含Gemini分析结果的字典
    """
    if not GEMINI_ENABLED:
        return {"error": "Gemini API not available"}

    try:
        # 确定使用的模型
        if model_key and model_key in GEMINI_MODELS:
            model_name = GEMINI_MODELS[model_key]["name"]
            logger.info(f"Using specified model: {model_key} -> {model_name}")
        else:
            model_name = DEFAULT_GEMINI_MODEL
            model_key = DEFAULT_GEMINI_MODEL_KEY
            logger.info(f"Using default model: {model_key} -> {model_name}")

        # 截断文本以避免超长输入
        truncated_text = truncate_text(text)
        logger.info(f"Text length: {len(text)} -> {len(truncated_text)} characters")

        # 获取合适的提示词
        prompt = get_prompt_for_model(model_key, language)

        # 创建完整的输入
        full_input = f"{prompt}\n\n文档内容：\n{truncated_text}"

        # 使用Gemini处理文本
        model = genai.GenerativeModel(model_name)

        logger.info(f"Starting Gemini analysis with model {model_name}")
        start_time = datetime.now()

        response = model.generate_content(
            full_input,
            generation_config={
                "temperature": 0.2,
                "top_p": 0.8,
                "top_k": 40,
                "max_output_tokens": 4096,  # 减少输出长度以提升速度
            }
        )

        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        logger.info(f"Gemini analysis completed in {processing_time:.2f} seconds")

        # 解析响应
        result = {
            "processed_at": datetime.now().isoformat(),
            "model_used": model_name,
            "model_key": model_key,
            "processing_time": processing_time,
            "text_length": len(truncated_text),
            "gemini_response": response.text
        }

        # 尝试解析JSON响应
        try:
            # 查找JSON部分
            json_start = response.text.find('{')
            json_end = response.text.rfind('}') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = response.text[json_start:json_end]
                parsed_json = json.loads(json_str)
                result["parsed_result"] = parsed_json

                # 提取关键信息
                for key in ["title", "document_type", "language", "summary", "key_points", "translation", "key_fields"]:
                    if key in parsed_json:
                        result[key] = parsed_json[key]

        except Exception as e:
            logger.warning(f"Failed to parse JSON from Gemini response: {e}")
            result["parsed_result"] = None

        return result

    except Exception as e:
        logger.exception(f"Error analyzing text with Gemini: {e}")
        return {"error": f"Error analyzing text with Gemini: {str(e)}"}

async def analyze_text_with_gemini(text: str, model_key: str = None, language: str = "auto") -> Dict[str, Any]:
    """
    使用Gemini分析文本内容（优化版本，只传输文本）

    Args:
        text: 要分析的文本内容
        model_key: 使用的模型键值，如果为None则使用默认模型
        language: 文档语言

    Returns:
        包含Gemini分析结果的字典
    """
    if not GEMINI_ENABLED:
        return {"error": "Gemini API not available"}

    try:
        # 确定使用的模型
        if model_key and model_key in GEMINI_MODELS:
            model_name = GEMINI_MODELS[model_key]["name"]
            logger.info(f"Using specified model: {model_key} -> {model_name}")
        else:
            model_name = DEFAULT_GEMINI_MODEL
            model_key = DEFAULT_GEMINI_MODEL_KEY
            logger.info(f"Using default model: {model_key} -> {model_name}")

        # 截断文本以避免超长输入
        truncated_text = truncate_text(text)
        logger.info(f"Text length: {len(text)} -> {len(truncated_text)} characters")

        # 获取合适的提示词
        prompt = get_prompt_for_model(model_key, language)

        # 创建完整的输入
        full_input = f"{prompt}\n\n文档内容：\n{truncated_text}"

        # 使用Gemini处理文本
        model = genai.GenerativeModel(model_name)

        logger.info(f"Starting Gemini analysis with model {model_name}")
        start_time = datetime.now()

        response = model.generate_content(
            full_input,
            generation_config={
                "temperature": 0.2,
                "top_p": 0.8,
                "top_k": 40,
                "max_output_tokens": 4096,  # 减少输出长度以提升速度
            }
        )

        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        logger.info(f"Gemini analysis completed in {processing_time:.2f} seconds")

        # 解析响应
        result = {
            "processed_at": datetime.now().isoformat(),
            "model_used": model_name,
            "model_key": model_key,
            "processing_time": processing_time,
            "text_length": len(truncated_text),
            "gemini_response": response.text
        }

        # 尝试解析JSON响应
        try:
            # 查找JSON部分
            json_start = response.text.find('{')
            json_end = response.text.rfind('}') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = response.text[json_start:json_end]
                parsed_json = json.loads(json_str)
                result["parsed_result"] = parsed_json

                # 提取关键信息
                for key in ["title", "document_type", "language", "summary", "key_points", "translation", "key_fields"]:
                    if key in parsed_json:
                        result[key] = parsed_json[key]

        except Exception as e:
            logger.warning(f"Failed to parse JSON from Gemini response: {e}")
            result["parsed_result"] = None

        return result

    except Exception as e:
        logger.exception(f"Error analyzing text with Gemini: {e}")
        return {"error": f"Error analyzing text with Gemini: {str(e)}"}

def process_with_gemini(pdf_path: str, prompt: Optional[str] = None, language: str = "auto") -> Dict[str, Any]:
    """
    使用Gemini处理PDF文件

    Args:
        pdf_path: PDF文件路径
        prompt: 可选的提示词，用于指导Gemini的处理
        language: 文档语言，用于选择合适的提示词

    Returns:
        包含Gemini处理结果的字典
    """
    if not GEMINI_ENABLED:
        return {"error": "Gemini API not available"}

    try:
        # 检查文件是否存在
        if not os.path.exists(pdf_path):
            return {"error": f"File not found: {pdf_path}"}

        # 读取PDF文件为bytes
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        # 将PDF转换为base64
        pdf_base64 = base64.b64encode(pdf_bytes).decode("utf-8")

        # 注意：不再使用 Part 类，因为它在当前版本中不可用

        # 设置默认提示词
        if not prompt:
            if language.lower() == "ja" or language.lower() == "japanese":
                prompt = DEFAULT_JAPANESE_PDF_PROMPT
            else:
                prompt = DEFAULT_PDF_PROMPT

        # 使用Gemini处理PDF
        model = genai.GenerativeModel(DEFAULT_GEMINI_MODEL)

        # 创建请求，尝试多种格式
        try:
            # 尝试方法1: 使用内联数据格式
            logger.info(f"Attempting to generate content with model {DEFAULT_GEMINI_MODEL} using inline data format.")

            # 创建请求内容
            content = [
                {"role": "user", "parts": [{"text": prompt}]},
                {"role": "user", "parts": [{"inline_data": {"mime_type": "application/pdf", "data": pdf_base64}}]}
            ]

            response = model.generate_content(
                content,
                generation_config={
                    "temperature": 0.2,
                    "top_p": 0.8,
                    "top_k": 40,
                    "max_output_tokens": 8192,
                }
            )
            logger.info("Successfully used inline data format with Gemini.")
        except Exception as e_inline_data:
            logger.warning(f"Failed to use inline data format: {e_inline_data}. Trying fallback.")

            # 尝试方法2: 使用简化的API格式
            try:
                logger.info("Trying simplified API format...")

                # 创建简化的请求内容
                response = model.generate_content(
                    prompt + "\n\n[PDF文档内容已提取，但由于技术限制无法直接传递]",
                    generation_config={
                        "temperature": 0.2,
                        "top_p": 0.8,
                        "top_k": 40,
                        "max_output_tokens": 8192,
                    }
                )
                logger.info("Successfully used simplified API format")
            except Exception as e_new_api:
                logger.warning(f"Failed to use new API format: {e_new_api}. Trying simpler approach.")

                # 尝试方法3: 使用最简单的方法，只发送提示文本
                try:
                    logger.info("Trying text-only approach with PDF content summary...")
                    # 尝试从PDF中提取一些文本作为上下文
                    context = ""
                    # 直接从文件中提取一些文本
                    if PYPDF2_AVAILABLE:
                        try:
                            with open(pdf_path, "rb") as f:
                                pdf_reader = PyPDF2.PdfReader(f)
                                if len(pdf_reader.pages) > 0:
                                    # 提取第一页文本
                                    context = pdf_reader.pages[0].extract_text() or ""
                                    # 限制长度
                                    context = context[:1000]
                        except Exception as e_extract:
                            logger.warning(f"Failed to extract text from PDF: {e_extract}")

                    # 创建带有上下文的提示
                    enhanced_prompt = f"{prompt}\n\n文档内容摘要:\n{context}"

                    response = model.generate_content(
                        enhanced_prompt,
                        generation_config={
                            "temperature": 0.2,
                            "top_p": 0.8,
                            "top_k": 40,
                            "max_output_tokens": 8192,
                        }
                    )
                    logger.info("Successfully used text-only approach")
                except Exception as e_text_only:
                    logger.error(f"All attempts to generate content with Gemini failed. Last error: {e_text_only}")
                    raise Exception(f"Failed to generate content with Gemini after trying all formats")

        # 解析响应
        result = {
            "processed_at": datetime.now().isoformat(),
            "original_filename": os.path.basename(pdf_path),
            "gemini_response": response.text
        }

        # 尝试解析JSON响应
        try:
            # 查找JSON部分
            json_start = response.text.find('{')
            json_end = response.text.rfind('}') + 1
            if json_start >= 0 and json_end > json_start:
                json_str = response.text[json_start:json_end]
                parsed_json = json.loads(json_str)
                result["parsed_result"] = parsed_json

                # 提取关键信息
                if "title" in parsed_json:
                    result["title"] = parsed_json["title"]
                if "document_type" in parsed_json:
                    result["document_type"] = parsed_json["document_type"]
                if "language" in parsed_json:
                    result["language"] = parsed_json["language"]
                if "summary" in parsed_json:
                    result["summary"] = parsed_json["summary"]
                if "key_points" in parsed_json:
                    result["key_points"] = parsed_json["key_points"]
                if "translation" in parsed_json:
                    translation_content = parsed_json["translation"]
                    if isinstance(translation_content, dict):
                        # 将字典的值拼接起来，或者根据需要选择特定字段
                        result["translation"] = " | ".join(str(v) for v in translation_content.values())
                        logger.info(f"Converted dict translation to string: {result['translation'][:200]}...")
                    elif isinstance(translation_content, str):
                        result["translation"] = translation_content
                    else:
                        result["translation"] = str(translation_content) # Fallback
        except Exception as e:
            logger.warning(f"Failed to parse JSON from Gemini response: {e}")
            result["parsed_result"] = None

        return result

    except Exception as e:
        logger.exception(f"Error processing file with Gemini: {e}")
        return {"error": f"Error processing file with Gemini: {str(e)}"}

def process_word(file_path: str, language: str = "auto", gemini_model: str = "gemini-1.5-flash") -> Dict[str, Any]:
    """
    处理Word文档，提取文本内容和表格数据

    Args:
        file_path: Word文档文件路径
        language: 语言代码
        gemini_model: Gemini模型选择

    Returns:
        包含处理结果的字典
    """
    start_time = datetime.now()
    logger.info(f"开始处理Word文档: {file_path}")

    try:
        if not WORD_PROCESSING_AVAILABLE:
            return {"error": "Word processing libraries not available"}

        # 读取Word文档
        doc = docx.Document(file_path)

        # 提取文本内容
        full_text_parts = []
        tables_data = []

        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text_parts.append(paragraph.text.strip())

        # 提取表格数据
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)

            if table_data:
                tables_data.append({
                    "table_id": f"table_{table_idx + 1}",
                    "rows": table_data,
                    "raw_text": "\n".join(["\t".join(row) for row in table_data])
                })

                # 将表格内容也添加到全文中
                full_text_parts.append(f"\n=== 表格 {table_idx + 1} ===")
                for row in table_data:
                    full_text_parts.append("\t".join(row))

        full_text = "\n".join(full_text_parts)

        # 创建页面内容（Word文档作为单页处理）
        pages = [{
            "page_number": 1,
            "text": full_text,
            "tables": [],
            "images": []
        }]

        # 处理信息
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()

        processing_info = {
            "pypdf2_used": False,
            "docling_used": False,
            "gemini_used": False,
            "force_ocr_used": False,
            "processing_time_seconds": processing_time,
            "status": "success",
            "error_message": None
        }

        # 文档元数据
        document_metadata = {
            "original_filename": os.path.basename(file_path),
            "processed_at": datetime.now().isoformat(),
            "page_count": 1,
            "source_format": "word",
            "language": language
        }

        result = {
            "document_metadata": document_metadata,
            "processing_info": processing_info,
            "pages": pages,
            "full_text": full_text,
            "tables": tables_data,
            "images": []
        }

        logger.info(f"Word文档处理完成: {file_path}, 提取文本长度: {len(full_text)}")
        return result

    except Exception as e:
        logger.exception(f"处理Word文档时发生错误: {file_path}")
        return {"error": f"处理Word文档失败: {str(e)}"}

def process_text_file(file_path: str, file_ext: str, language: str = "auto") -> Dict[str, Any]:
    """
    处理文本文件（TXT、MD、RTF、CSV、TSV）

    Args:
        file_path: 文本文件路径
        file_ext: 文件扩展名
        language: 语言代码

    Returns:
        包含处理结果的字典
    """
    start_time = datetime.now()
    logger.info(f"开始处理文本文件: {file_path}, 类型: {file_ext}")

    try:
        # 尝试不同编码读取文件
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        content = None
        used_encoding = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                logger.info(f"成功使用 {encoding} 编码读取文件")
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            return {"error": "无法读取文件，尝试了多种编码格式"}

        # 根据文件类型处理内容
        if file_ext in ['.csv', '.tsv']:
            # CSV/TSV文件特殊处理
            lines = content.strip().split('\n')
            delimiter = '\t' if file_ext == '.tsv' else ','

            processed_content = f"=== {file_ext.upper()} 数据 ===\n"
            tables_data = []

            if lines:
                # 解析CSV/TSV数据
                table_rows = []
                for line in lines:
                    if line.strip():
                        fields = line.split(delimiter)
                        table_rows.append([field.strip() for field in fields])

                if table_rows:
                    tables_data.append({
                        "table_id": "csv_table_1",
                        "rows": table_rows,
                        "raw_text": content
                    })

                    # 格式化显示
                    for i, row in enumerate(table_rows[:50]):  # 限制显示50行
                        if i == 0:
                            processed_content += f"列标题: {' | '.join(row)}\n"
                        else:
                            processed_content += f"第{i}行: {' | '.join(row)}\n"

            full_text = processed_content
        else:
            # 普通文本文件
            full_text = content
            tables_data = []

        # 创建页面内容
        pages = [{
            "page_number": 1,
            "text": full_text,
            "tables": [],
            "images": []
        }]

        # 处理信息
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()

        processing_info = {
            "pypdf2_used": False,
            "docling_used": False,
            "gemini_used": False,
            "force_ocr_used": False,
            "processing_time_seconds": processing_time,
            "status": "success",
            "error_message": None
        }

        # 文档元数据
        document_metadata = {
            "original_filename": os.path.basename(file_path),
            "processed_at": datetime.now().isoformat(),
            "page_count": 1,
            "source_format": f"text{file_ext}",
            "language": language,
            "encoding": used_encoding
        }

        result = {
            "document_metadata": document_metadata,
            "processing_info": processing_info,
            "pages": pages,
            "full_text": full_text,
            "tables": tables_data,
            "images": []
        }

        logger.info(f"文本文件处理完成: {file_path}, 提取文本长度: {len(full_text)}")
        return result

    except Exception as e:
        logger.exception(f"处理文本文件时发生错误: {file_path}")
        return {"error": f"处理文本文件失败: {str(e)}"}

def process_excel(file_path: str, language: str = "auto", gemini_model: str = "gemini-1.5-flash") -> Dict[str, Any]:
    """
    处理Excel文件，提取文本内容和表格数据

    Args:
        file_path: Excel文件路径
        language: 文档语言
        gemini_model: Gemini模型选择

    Returns:
        包含处理结果的字典
    """
    start_time = datetime.now()

    if not EXCEL_PROCESSING_AVAILABLE:
        return {"error": "Excel processing libraries not available"}

    try:
        logger.info(f"Processing Excel file: {file_path}")

        # 检测文件扩展名
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in ['.xlsx', '.xls', '.xlsm']:
            return {"error": f"Unsupported Excel file format: {file_ext}"}

        # 读取Excel文件
        try:
            if file_ext == '.xls':
                # 对于旧版Excel文件，使用xlrd引擎
                df_dict = pd.read_excel(file_path, sheet_name=None, engine='xlrd')
            else:
                # 对于新版Excel文件，使用openpyxl引擎
                df_dict = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
        except Exception as e:
            logger.error(f"Failed to read Excel file: {e}")
            return {"error": f"Failed to read Excel file: {str(e)}"}

        # 处理结果
        result = {
            "document_metadata": {
                "original_filename": os.path.basename(file_path),
                "processed_at": datetime.now().isoformat(),
                "source_format": "Excel",
                "sheet_count": len(df_dict)
            },
            "processing_info": {
                "pypdf2_used": False,
                "docling_used": False,
                "gemini_used": False,
                "force_ocr_used": False,
                "status": "success"
            },
            "pages": [],
            "full_text": "",
            "tables": [],
            "images": []
        }

        all_text_parts = []
        sheet_number = 1

        # 处理每个工作表
        for sheet_name, df in df_dict.items():
            logger.info(f"Processing sheet: {sheet_name}")

            # 构建工作表文本
            sheet_text_parts = [f"=== 工作表: {sheet_name} ==="]

            # 检查是否有数据
            if df.empty:
                sheet_text_parts.append("(空工作表)")
            else:
                # 添加列标题
                if not df.columns.empty:
                    headers = [str(col) for col in df.columns]
                    sheet_text_parts.append("列标题: " + " | ".join(headers))

                # 添加数据行（限制行数以避免过长）
                max_rows = 100  # 限制最多处理100行
                for idx, row in df.head(max_rows).iterrows():
                    row_values = []
                    for val in row:
                        if pd.isna(val):
                            row_values.append("")
                        else:
                            row_values.append(str(val))
                    sheet_text_parts.append(" | ".join(row_values))

                # 如果行数超过限制，添加提示
                if len(df) > max_rows:
                    sheet_text_parts.append(f"... (还有 {len(df) - max_rows} 行数据)")

            sheet_text = "\n".join(sheet_text_parts)
            all_text_parts.append(sheet_text)

            # 创建页面内容（每个工作表作为一页）
            page_content = {
                "page_number": sheet_number,
                "text": sheet_text,
                "tables": [],
                "images": []
            }

            # 如果工作表有数据，创建表格信息
            if not df.empty:
                table_info = {
                    "table_id": f"sheet_{sheet_number}_{sheet_name}",
                    "page_number": sheet_number,
                    "title": f"工作表: {sheet_name}",
                    "headers": [str(col) for col in df.columns] if not df.columns.empty else [],
                    "rows": [],
                    "raw_text": sheet_text
                }

                # 添加数据行（限制行数）
                for idx, row in df.head(50).iterrows():  # 表格数据限制50行
                    row_values = []
                    for val in row:
                        if pd.isna(val):
                            row_values.append("")
                        else:
                            row_values.append(str(val))
                    table_info["rows"].append(row_values)

                page_content["tables"].append(table_info)
                result["tables"].append(table_info)

            result["pages"].append(page_content)
            sheet_number += 1

        # 合并所有文本
        result["full_text"] = "\n\n".join(all_text_parts)

        # 计算处理时间
        end_time = datetime.now()
        processing_time = (end_time - start_time).total_seconds()
        result["processing_info"]["processing_time_seconds"] = processing_time

        logger.info(f"Excel processing completed in {processing_time:.2f} seconds")
        logger.info(f"Extracted text length: {len(result['full_text'])} characters")

        return result

    except Exception as e:
        logger.exception(f"Error processing Excel file: {e}")
        return {"error": f"Error processing Excel file: {str(e)}"}

def process_pdf(file_path: str, use_pypdf2: bool = True, use_docling: bool = True,
                use_gemini: bool = True, force_ocr: bool = False,
                gemini_prompt: Optional[str] = None, language: str = "auto",
                gemini_model: str = "gemini-1.5-flash") -> Dict[str, Any]:
    """
    综合处理PDF文件，结合PyPDF2、Docling和Gemini的能力

    Args:
        file_path: PDF文件路径
        use_pypdf2: 是否使用PyPDF2处理
        use_docling: 是否使用Docling处理
        use_gemini: 是否使用Gemini处理
        force_ocr: 是否强制使用OCR（仅适用于Docling）
        gemini_prompt: 可选的Gemini提示词
        language: 文档语言，用于选择合适的提示词
        gemini_model: Gemini模型选择

    Returns:
        包含处理结果的字典
    """
    start_time = datetime.now()

    # 创建处理信息
    processing_info = {
        "pypdf2_used": use_pypdf2 and PYPDF2_AVAILABLE,
        "docling_used": use_docling and DOCLING_AVAILABLE,
        "gemini_used": use_gemini and GEMINI_ENABLED,
        "force_ocr_used": force_ocr,
        "status": "success"
    }

    # 创建结果字典
    result = {
        "document_metadata": {
            "original_filename": os.path.basename(file_path),
            "processed_at": datetime.now().isoformat()
        },
        "processing_info": processing_info,
        "pages": [],
        "full_text": "",
        "tables": [],
        "images": []
    }

    # 使用PyPDF2处理
    pypdf2_result = None
    if use_pypdf2 and PYPDF2_AVAILABLE:
        logger.info(f"Processing with PyPDF2: {file_path}")
        pypdf2_result = extract_text_with_pypdf2(file_path)

        if "error" not in pypdf2_result:
            # 更新文档元数据
            if "page_count" in pypdf2_result:
                result["document_metadata"]["page_count"] = pypdf2_result["page_count"]
            if "title" in pypdf2_result:
                result["document_metadata"]["title"] = pypdf2_result["title"]
            if "author" in pypdf2_result:
                result["document_metadata"]["author"] = pypdf2_result["author"]
            if "creation_date" in pypdf2_result:
                result["document_metadata"]["creation_date"] = pypdf2_result["creation_date"]

            # 更新页面内容
            if "pages" in pypdf2_result:
                for page_data in pypdf2_result["pages"]:
                    result["pages"].append({
                        "page_number": page_data["page_number"],
                        "text": page_data["text"],
                        "tables": [],
                        "images": []
                    })

            # 更新全文
            if "full_text" in pypdf2_result:
                result["full_text"] = pypdf2_result["full_text"]
        else:
            logger.warning(f"PyPDF2 processing failed: {pypdf2_result['error']}")

    # 提取PDF中的图像（使用PyMuPDF）- 独立于其他处理选项
    logger.info(f"Extracting images from PDF: {file_path}")
    try:
        extracted_images = extract_images_from_pdf(file_path)
        if extracted_images:
            result["images"].extend(extracted_images)
            logger.info(f"Successfully extracted {len(extracted_images)} images from PDF")

            # 将图像信息添加到对应的页面中，并在文本中插入图像标记
            for image_info in extracted_images:
                page_number = image_info.page_number
                # 查找对应的页面并添加图像信息
                for page in result["pages"]:
                    if page["page_number"] == page_number:
                        page["images"].append({
                            "image_id": image_info.image_id,
                            "page_number": image_info.page_number,
                            "description": image_info.description,
                            "mime_type": image_info.mime_type,
                            "data": image_info.data
                        })

                        # 在页面文本中插入图像标记
                        image_marker = f"\n[IMAGE:{image_info.image_id}:{image_info.description}]\n"
                        if page["text"]:
                            # 如果页面有文本，在文本开头插入图像标记
                            page["text"] = image_marker + page["text"]
                        else:
                            # 如果页面没有文本，只插入图像标记
                            page["text"] = image_marker
                        break
                # 如果没有找到对应页面，创建一个新页面（防止页面信息缺失）
                if not any(page["page_number"] == page_number for page in result["pages"]):
                    image_marker = f"\n[IMAGE:{image_info.image_id}:{image_info.description}]\n"
                    result["pages"].append({
                        "page_number": page_number,
                        "text": image_marker,
                        "tables": [],
                        "images": [{
                            "image_id": image_info.image_id,
                            "page_number": image_info.page_number,
                            "description": image_info.description,
                            "mime_type": image_info.mime_type,
                            "data": image_info.data
                        }]
                    })

            # 重新构建全文，包含图像标记
            updated_full_text = ""
            for page in sorted(result["pages"], key=lambda x: x["page_number"]):
                if page["text"]:
                    updated_full_text += page["text"] + "\n"
            result["full_text"] = updated_full_text.strip()

        else:
            logger.info("No images found in PDF")
    except Exception as e:
        logger.error(f"Failed to extract images from PDF: {e}")

    # 使用Docling处理
    docling_result = None
    if use_docling and DOCLING_AVAILABLE:
        logger.info(f"Processing with Docling: {file_path}")
        docling_result = process_with_docling(file_path, force_ocr)

        if "error" not in docling_result:
            # 如果PyPDF2没有提取到文本，使用Docling的结果
            if not result["full_text"] and "full_text" in docling_result:
                result["full_text"] = docling_result["full_text"]

            # 如果PyPDF2没有提取到页面内容，使用Docling的结果
            if not result["pages"] and "pages" in docling_result:
                for page_data in docling_result["pages"]:
                    result["pages"].append({
                        "page_number": page_data["page_number"],
                        "text": page_data.get("text", ""),
                        "tables": [],
                        "images": []
                    })

            # 更新文档元数据
            if "page_count" in docling_result and not result["document_metadata"].get("page_count"):
                result["document_metadata"]["page_count"] = docling_result["page_count"]
        else:
            logger.warning(f"Docling processing failed: {docling_result['error']}")

    # 使用Gemini处理（优化版本：只传输文本内容）
    gemini_result = None
    if use_gemini and GEMINI_ENABLED:
        logger.info(f"Processing with Gemini (optimized): {file_path}")

        # 获取已提取的文本内容
        extracted_text = result["full_text"]
        if extracted_text:
            # 使用优化的文本分析函数（同步版本）
            gemini_result = analyze_text_with_gemini_sync(
                text=extracted_text,
                model_key=gemini_model,
                language=language
            )
        else:
            logger.warning("No text extracted for Gemini analysis, skipping Gemini processing")
            gemini_result = {"error": "No text available for analysis"}

        if "error" not in gemini_result:
            # 创建Gemini分析结果
            gemini_analysis = {
                "raw_response": gemini_result.get("gemini_response", "")
            }

            # 提取解析结果
            if "parsed_result" in gemini_result:
                parsed_result = gemini_result["parsed_result"]

                # 提取摘要
                if "summary" in parsed_result:
                    gemini_analysis["summary"] = parsed_result["summary"]

                # 提取关键点
                if "key_points" in parsed_result:
                    gemini_analysis["key_points"] = parsed_result["key_points"]

                # 提取结构化数据
                gemini_analysis["structured_data"] = parsed_result

                # 提取翻译
                if "translation" in parsed_result:
                    translation_content = parsed_result["translation"]
                    if isinstance(translation_content, dict):
                        # 将字典的值拼接起来，或者根据需要选择特定字段
                        gemini_analysis["translation"] = " | ".join(str(v) for v in translation_content.values())
                        logger.info(f"Converted dict translation to string: {gemini_analysis['translation'][:200]}...")
                    elif isinstance(translation_content, str):
                        gemini_analysis["translation"] = translation_content
                    else:
                        gemini_analysis["translation"] = str(translation_content) # Fallback

                # 提取表格
                if "tables" in parsed_result:
                    tables = parsed_result["tables"]
                    for i, table in enumerate(tables):
                        try:
                            table_info = {
                                "table_id": f"table_{i+1}",
                                "page_number": 1,  # 默认值，Gemini可能无法确定表格所在页面
                                "title": table.get("table_title"),
                                "raw_text": table.get("content"),
                                "rows": []  # 初始化为空列表
                            }

                            # 尝试提取表格行和列
                            content = table.get("content")
                            if isinstance(content, dict) and "headers" in content and "rows" in content:
                                table_info["headers"] = content["headers"]
                                # 确保rows是列表的列表
                                if isinstance(content["rows"], list):
                                    if content["rows"] and isinstance(content["rows"][0], list):
                                        # 已经是正确的格式
                                        table_info["rows"] = content["rows"]
                                    elif content["rows"] and isinstance(content["rows"][0], dict):
                                        # 字典列表，需要转换
                                        rows = []
                                        for row_dict in content["rows"]:
                                            row = [str(row_dict.get(header, "")) for header in content["headers"]]
                                            rows.append(row)
                                        table_info["rows"] = rows
                                    else:
                                        # 其他情况，使用空列表
                                        table_info["rows"] = []
                                else:
                                    # 不是列表，使用空列表
                                    table_info["rows"] = []
                            elif isinstance(content, list) and all(isinstance(item, dict) for item in content):
                                # 如果内容是字典列表，尝试提取表头和行
                                if content:
                                    headers = list(content[0].keys())
                                    table_info["headers"] = headers
                                    rows = []
                                    for item in content:
                                        row = [str(item.get(header, "")) for header in headers]
                                        rows.append(row)
                                    table_info["rows"] = rows
                            elif isinstance(content, str):
                                # 如果内容是字符串，尝试解析表格结构
                                try:
                                    # 分割行
                                    lines = content.strip().split('\n')
                                    if lines:
                                        # 第一行可能是表头
                                        header_line = lines[0]
                                        # 使用 | 分割列
                                        if '|' in header_line:
                                            headers = [h.strip() for h in header_line.split('|')]
                                            table_info["headers"] = headers

                                            # 处理数据行
                                            rows = []
                                            for line in lines[1:]:
                                                if '|' in line:
                                                    row = [cell.strip() for cell in line.split('|')]
                                                    rows.append(row)
                                            table_info["rows"] = rows
                                except Exception as e:
                                    logger.warning(f"Failed to parse table content string: {e}")
                                    # 保存原始文本
                                    table_info["raw_text"] = content

                            # 确保rows是列表的列表
                            if "rows" in table_info and not isinstance(table_info["rows"], list):
                                table_info["rows"] = []
                            elif "rows" in table_info and table_info["rows"] and not isinstance(table_info["rows"][0], list):
                                # 如果rows是字典列表，尝试转换
                                if isinstance(table_info["rows"], list) and table_info["rows"] and isinstance(table_info["rows"][0], dict):
                                    # 获取所有键
                                    all_keys = set()
                                    for row_dict in table_info["rows"]:
                                        all_keys.update(row_dict.keys())

                                    # 转换为有序列表
                                    keys_list = sorted(list(all_keys))

                                    # 创建表头
                                    if "headers" not in table_info or not table_info["headers"]:
                                        table_info["headers"] = keys_list

                                    # 转换每一行
                                    new_rows = []
                                    for row_dict in table_info["rows"]:
                                        row_values = [str(row_dict.get(key, "")) for key in keys_list]
                                        new_rows.append(row_values)

                                    table_info["rows"] = new_rows
                                else:
                                    # 其他情况，使用空列表
                                    table_info["rows"] = []

                            result["tables"].append(table_info)
                        except Exception as e:
                            logger.warning(f"Failed to process table {i+1}: {e}")
                            # 添加一个简单的表格信息
                            result["tables"].append({
                                "table_id": f"table_{i+1}",
                                "page_number": 1,
                                "title": table.get("table_title", f"Table {i+1}"),
                                "raw_text": str(table.get("content", "")),
                                "rows": []
                            })

                # 提取图像
                if "images" in parsed_result:
                    images = parsed_result["images"]
                    for i, image in enumerate(images):
                        image_info = {
                            "image_id": f"image_{i+1}",
                            "page_number": 1,  # 默认值，Gemini可能无法确定图像所在页面
                            "description": image.get("description") or image.get("image_description")
                        }
                        result["images"].append(image_info)

                # 更新文档元数据
                if "title" in parsed_result and not result["document_metadata"].get("title"):
                    result["document_metadata"]["title"] = parsed_result["title"]
                if "document_type" in parsed_result:
                    result["document_metadata"]["document_type"] = parsed_result["document_type"]
                if "language" in parsed_result:
                    result["document_metadata"]["language"] = parsed_result["language"]

            result["gemini_analysis"] = gemini_analysis
        else:
            logger.warning(f"Gemini processing failed: {gemini_result['error']}")

    # 计算处理时间
    end_time = datetime.now()
    processing_time = (end_time - start_time).total_seconds()
    result["processing_info"]["processing_time_seconds"] = processing_time

    # 检查处理结果
    if not result["full_text"] and not result["pages"]:
        result["processing_info"]["status"] = "error"
        result["processing_info"]["error_message"] = "Failed to extract any text from the document"

    return result


async def process_with_gemini_vision_ocr(file_path: str, language: str, model_name: str) -> Dict[str, Any]:
    """
    使用Gemini Vision OCR处理各种文件类型（PDF、图片、文档等）。

    Args:
        file_path: 文件路径（支持PDF、图片、文档等）。
        language: 提取语言 (传递给Gemini的prompt，可以是 'auto')。
        model_name: 使用的Gemini模型名称。

    Returns:
        包含提取文本和页码的字典。
    """
    if not GEMINI_ENABLED:
        return {"error": "Gemini AI not enabled or API key not configured."}

    # 检查文件类型，智能处理不同格式
    file_ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"Gemini Vision OCR processing file: {file_path}, extension: {file_ext}")

    all_text: List[str] = []
    images_to_process = []

    try:
        # 根据文件类型准备图像数据
        if file_ext == '.pdf':
            # PDF文件：使用PyMuPDF转换为图像
            with open(file_path, "rb") as f:
                pdf_content_bytes = f.read()

            doc = fitz.open(stream=pdf_content_bytes, filetype="pdf")
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(dpi=300)
                img_bytes = pix.tobytes("png")
                img_pil = Image.open(io.BytesIO(img_bytes))
                images_to_process.append({
                    "image": img_pil,
                    "page_number": page_num + 1,
                    "source": f"PDF page {page_num + 1}"
                })
            doc.close()

        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp']:
            # 图片文件：直接处理
            img_pil = Image.open(file_path)
            images_to_process.append({
                "image": img_pil,
                "page_number": 1,
                "source": "Image file"
            })

        elif file_ext in ['.txt', '.md', '.rtf', '.csv', '.tsv']:
            # 文本文件：读取内容并创建文本图像（可选）或直接返回文本
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()

            # 对于文本文件，我们可以直接返回内容，不需要OCR
            logger.info(f"Text file detected, returning content directly (length: {len(text_content)})")
            return {
                "extracted_text": text_content,
                "page_count": 1,
                "warning": "Text file processed directly without OCR"
            }

        elif file_ext in ['.docx', '.doc', '.xlsx', '.xls', '.xlsm']:
            # Office文档：提示用户这些文件应该用专门的处理器
            return {
                "error": f"Office documents ({file_ext}) should be processed with dedicated handlers, not Vision OCR"
            }

        else:
            # 未知文件类型：尝试作为图像处理
            try:
                img_pil = Image.open(file_path)
                images_to_process.append({
                    "image": img_pil,
                    "page_number": 1,
                    "source": f"Unknown file type {file_ext} treated as image"
                })
                logger.info(f"Unknown file type {file_ext} successfully opened as image")
            except Exception as e:
                return {"error": f"Unsupported file type {file_ext} and cannot open as image: {str(e)}"}

        # 处理准备好的图像
        if not images_to_process:
            return {"error": "No images to process"}

        model = genai.GenerativeModel(model_name)

        # 基于用户当前的极简提示词进行修改
        final_minimal_prompt = f"""Extract all text from this image. Document language is {language if language != 'auto' else 'any language'}. Output in the original format."""

        for img_data in images_to_process:
            img_pil = img_data["image"]
            page_number = img_data["page_number"]
            source = img_data["source"]

            logger.info(f"向Gemini发送{source}进行OCR处理...")

            prompt_parts_for_api = [final_minimal_prompt] # 使用最终的极简提示词
            
            current_page_text_parts = []
            try:
                # 保持120秒超时，观察效果
                response = model.generate_content(
                    prompt_parts_for_api + [img_pil],
                    request_options={"timeout": 120} 
                )

                if not response.candidates:
                    logger.warning(f"第 {page_num + 1} 页 Gemini 响应中没有候选内容 (no candidates)。")
                    continue # 跳过此页

                candidate = response.candidates[0]
                
                # candidate.finish_reason 是 google.generativeai.types.Candidate.FinishReason 枚举
                # .value: STOP=1, MAX_TOKENS=2, SAFETY=3, RECITATION=4, OTHER=5
                
                if candidate.finish_reason.value == 1: # STOP - 正常完成
                    if candidate.content and candidate.content.parts:
                        for part in candidate.content.parts:
                            if hasattr(part, "text") and part.text:
                                current_page_text_parts.append(part.text)
                        
                        if current_page_text_parts:
                            page_text_content = "".join(current_page_text_parts)
                            all_text.append(page_text_content)
                            logger.info(f"成功处理{source}，提取文本长度: {len(page_text_content)}")
                        else:
                            logger.warning(f"{source} Gemini 响应状态为 STOP，但未找到文本内容。 Parts: {[p.to_dict() for p in candidate.content.parts] if candidate.content else 'N/A'}")
                    else:
                        logger.warning(f"{source} Gemini 响应状态为 STOP，但 Candidate.content 或 Candidate.content.parts 为空。")
                
                elif candidate.finish_reason.value == 2: # MAX_TOKENS
                    logger.warning(
                        f"{source} Gemini 处理因 MAX_TOKENS (输出Token达到上限) 而结束。将尝试提取部分文本。"
                        f"Finish Reason: {candidate.finish_reason.name} (Value: {candidate.finish_reason.value}). "
                        f"Safety Ratings: {[str(rating) for rating in candidate.safety_ratings] if candidate.safety_ratings else 'N/A'}"
                    )
                    if candidate.content and candidate.content.parts:
                        for part in candidate.content.parts:
                            if hasattr(part, "text") and part.text:
                                current_page_text_parts.append(part.text)
                        if current_page_text_parts:
                            partial_text = "".join(current_page_text_parts)
                            all_text.append(partial_text) # 添加部分文本
                            logger.info(f"从{source}提取到部分文本 (MAX_TOKENS)，长度: {len(partial_text)}")
                        else:
                            logger.warning(f"{source}因 MAX_TOKENS 结束，但未在响应的 parts 中找到部分文本。")

                else: # 其他非正常结束原因 (SAFETY, RECITATION, OTHER, UNSPECIFIED)
                    logger.warning(
                        f"{source} Gemini 处理未正常完成。将跳过此页的文本提取。"
                        f"Finish Reason: {candidate.finish_reason.name} (Value: {candidate.finish_reason.value}). "
                        f"Safety Ratings: {[str(rating) for rating in candidate.safety_ratings] if candidate.safety_ratings else 'N/A'}"
                    )

            except AttributeError as ae:
                logger.error(f"处理{source}的Gemini响应时发生 AttributeError (API响应结构可能不符合预期): {ae}", exc_info=True)
                if 'response' in locals() and response:
                     logger.error(f"原始 Gemini 响应 ({source}): {response}")
            except Exception as e_page:
                logger.error(f"处理{source}时发生错误: {e_page}", exc_info=True)
        
        final_extracted_text = "\n\n".join(all_text) # 使用两个换行符分隔页面文本
        total_images = len(images_to_process)
        logger.info(f"Gemini Vision OCR 完成，总提取文本长度: {len(final_extracted_text)}, 总处理图像数: {total_images}")
        return {"extracted_text": final_extracted_text, "page_count": total_images, "warning": "使用Gemini Vision OCR提取的文本"}

    except RuntimeError as e_pymupdf: # Catch RuntimeError which PyMuPDF often uses for internal errors
        logger.error(f"PyMuPDF runtime error during Vision OCR: {e_pymupdf}", exc_info=True)
        # Check if specific MuPDF error strings are in the message if needed
        if "FZ_ERROR_GENERIC" in str(e_pymupdf):
            return {"error": f"PyMuPDF generic error: {e_pymupdf}"}
        elif "FZ_ERROR_TRYLATER" in str(e_pymupdf):
            return {"error": f"PyMuPDF try later error: {e_pymupdf}"}
        return {"error": f"PyMuPDF runtime error: {e_pymupdf}"} # Default return for other RuntimeErrors from PyMuPDF
    except Exception as e:
        logger.error(f"使用Gemini Vision OCR处理文件 '{file_path}' 时发生未知错误: {e}", exc_info=True)
        return {"error": f"An unexpected error occurred during Gemini Vision OCR processing: {e}"}
    finally:
        # 清理资源（如果有的话）
        pass


# --- API Endpoints ---
@app.post("/api/ocr/upload", response_model=OcrResponse)
async def ocr_upload(
    request: Request, # 新增 Request 对象
    file: UploadFile = File(...),
    use_pypdf2: str = Form("true"), # 从表单接收
    use_docling: str = Form("true"), # 从表单接收
    use_gemini: str = Form("true"),   # 从表单接收
    force_ocr: str = Form("false"),  # 从表单接收
    language: str = Form("auto"),    # 从表单接收
    gemini_model: str = Form("gemini-1.5-flash"),  # 从表单接收
    use_vision_ocr: str = Form("false")  # 从表单接收，这是关键参数
):
    """
    接收文件上传，使用多层处理策略进行OCR和文档结构解析，
    并返回结构化的JSON。

    Args:
        file: 上传的PDF文件
        use_pypdf2: 是否使用PyPDF2处理
        use_docling: 是否使用Docling处理
        use_gemini: 是否使用Gemini处理
        force_ocr: 是否强制使用OCR（仅适用于Docling）
        language: 文档语言，用于选择合适的提示词
        gemini_model: Gemini模型选择 (gemini-2.5-pro, gemini-1.5-pro, gemini-1.5-flash)

    Returns:
        OcrResponse: 包含处理结果的结构化JSON
    """
    temp_file_path = None
    try:
        # 记录原始表单数据
        try:
            form_data = await request.form()
            logger.info(f"Received /api/ocr/upload request. Raw form data: {form_data}")
            if "use_vision_ocr" in form_data:
                logger.info(f"Raw form data 'use_vision_ocr': {form_data['use_vision_ocr']}")
            else:
                logger.warning("'use_vision_ocr' not found in raw form data.")
            logger.info(f"Parsed parameters: file_name='{file.filename}', use_pypdf2='{use_pypdf2}', use_docling='{use_docling}', use_gemini='{use_gemini}', force_ocr='{force_ocr}', language='{language}', gemini_model='{gemini_model}', use_vision_ocr (parameter)='{use_vision_ocr}'")
        except Exception as e:
            logger.error(f"Error logging form data: {e}")

        # 将字符串参数转换为boolean
        use_pypdf2_bool = use_pypdf2.lower() in ("true", "1", "yes", "on")
        use_docling_bool = use_docling.lower() in ("true", "1", "yes", "on")
        use_gemini_bool = use_gemini.lower() in ("true", "1", "yes", "on")
        force_ocr_bool = force_ocr.lower() in ("true", "1", "yes", "on")
        use_vision_ocr_bool = use_vision_ocr.lower() in ("true", "1", "yes", "on")

        logger.info(f"转换后的boolean参数: use_pypdf2={use_pypdf2_bool}, use_docling={use_docling_bool}, use_gemini={use_gemini_bool}, force_ocr={force_ocr_bool}, use_vision_ocr={use_vision_ocr_bool}")

        safe_filename = "".join(c if c.isalnum() or c in ('.', '_', '-') else '_' for c in file.filename)
        temp_file_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4().hex}_{safe_filename}")

        with open(temp_file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        logger.info(f"File '{file.filename}' uploaded to '{temp_file_path}'")

        # 检测文件类型
        file_ext = os.path.splitext(file.filename)[1].lower()
        logger.info(f"Processing file: {temp_file_path}, file extension: {file_ext}")

        # 检查是否为Excel文件
        if file_ext in ['.xlsx', '.xls', '.xlsm']:
            logger.info("Detected Excel file, processing with Excel handler")
            if not EXCEL_PROCESSING_AVAILABLE:
                raise HTTPException(status_code=400, detail="Excel processing libraries not available")

            # 处理Excel文件
            excel_result = process_excel(temp_file_path, language, gemini_model)

            if "error" in excel_result:
                logger.error(f"Excel processing failed: {excel_result['error']}")
                raise HTTPException(status_code=500, detail=f"Excel processing failed: {excel_result['error']}")

            # 转换为OcrResponse格式
            processing_info = ProcessingInfo(**excel_result["processing_info"])
            document_metadata = DocumentMetadata(**excel_result["document_metadata"])

            pages = []
            for page_data in excel_result.get("pages", []):
                page = PageContent(
                    page_number=page_data["page_number"],
                    text=page_data.get("text", ""),
                    tables=page_data.get("tables", []),
                    images=page_data.get("images", [])
                )
                pages.append(page)

            tables = []
            for table_data in excel_result.get("tables", []):
                try:
                    tables.append(TableInfo(**table_data))
                except Exception as e_table:
                    logger.warning(f"Skipping table due to parsing error: {e_table}")

            return OcrResponse(
                document_metadata=document_metadata,
                processing_info=processing_info,
                pages=pages,
                full_text=excel_result.get("full_text", ""),
                tables=tables,
                images=excel_result.get("images", []),
                gemini_analysis=None
            )

        # 检查是否为Word文档
        elif file_ext in ['.docx', '.doc']:
            logger.info("Detected Word document, processing with Word handler")
            if not WORD_PROCESSING_AVAILABLE:
                raise HTTPException(status_code=400, detail="Word processing libraries not available")

            # 处理Word文档
            word_result = process_word(temp_file_path, language, gemini_model)

            if "error" in word_result:
                logger.error(f"Word processing failed: {word_result['error']}")
                raise HTTPException(status_code=500, detail=f"Word processing failed: {word_result['error']}")

            # 转换为OcrResponse格式
            processing_info = ProcessingInfo(**word_result["processing_info"])
            document_metadata = DocumentMetadata(**word_result["document_metadata"])

            pages = []
            for page_data in word_result.get("pages", []):
                page = PageContent(
                    page_number=page_data["page_number"],
                    text=page_data.get("text", ""),
                    tables=page_data.get("tables", []),
                    images=page_data.get("images", [])
                )
                pages.append(page)

            tables = []
            for table_data in word_result.get("tables", []):
                try:
                    tables.append(TableInfo(**table_data))
                except Exception as e_table:
                    logger.warning(f"Skipping table due to parsing error: {e_table}")

            return OcrResponse(
                document_metadata=document_metadata,
                processing_info=processing_info,
                pages=pages,
                full_text=word_result.get("full_text", ""),
                tables=tables,
                images=word_result.get("images", []),
                gemini_analysis=None
            )

        # 检查是否为文本文件
        elif file_ext in ['.txt', '.md', '.rtf', '.csv', '.tsv']:
            logger.info(f"Detected text file ({file_ext}), processing with text handler")

            # 处理文本文件
            text_result = process_text_file(temp_file_path, file_ext, language)

            if "error" in text_result:
                logger.error(f"Text file processing failed: {text_result['error']}")
                raise HTTPException(status_code=500, detail=f"Text file processing failed: {text_result['error']}")

            # 转换为OcrResponse格式
            processing_info = ProcessingInfo(**text_result["processing_info"])
            document_metadata = DocumentMetadata(**text_result["document_metadata"])

            pages = []
            for page_data in text_result.get("pages", []):
                page = PageContent(
                    page_number=page_data["page_number"],
                    text=page_data.get("text", ""),
                    tables=page_data.get("tables", []),
                    images=page_data.get("images", [])
                )
                pages.append(page)

            return OcrResponse(
                document_metadata=document_metadata,
                processing_info=processing_info,
                pages=pages,
                full_text=text_result.get("full_text", ""),
                tables=text_result.get("tables", []),
                images=text_result.get("images", []),
                gemini_analysis=None
            )

        # 如果到这里，说明文件不是Excel、Word或文本文件，应该是PDF或图片文件
        # 对于PDF文件，根据use_vision_ocr_bool决定处理方式
        if use_vision_ocr_bool:
            vision_start_time = datetime.now()
            result_vision = await process_with_gemini_vision_ocr(temp_file_path, language, gemini_model)
            vision_end_time = datetime.now()
            vision_processing_time = (vision_end_time - vision_start_time).total_seconds()

            if "error" in result_vision:
                logger.error(f"Gemini Vision OCR processing failed: {result_vision['error']}")
                error_processing_info = ProcessingInfo(
                    pypdf2_used=False,
                    docling_used=False,
                    gemini_used=True, # Attempted Gemini Vision
                    force_ocr_used=True, # Vision is inherently OCR
                    status="error",
                    error_message=str(result_vision["error"]),
                    processing_time_seconds=vision_processing_time
                )
                error_document_metadata = DocumentMetadata(
                    original_filename=safe_filename,
                    processed_at=datetime.now().isoformat()
                )
                return OcrResponse(
                    document_metadata=error_document_metadata,
                    processing_info=error_processing_info,
                    pages=[], full_text="", tables=[], images=[], gemini_analysis=None
                )
            else:
                logger.info("Gemini Vision OCR successful.")
                processing_info_vision = ProcessingInfo(
                    pypdf2_used=False,
                    docling_used=False,
                    gemini_used=True,
                    force_ocr_used=True,
                    processing_time_seconds=vision_processing_time,
                    status="success",
                    error_message=result_vision.get("warning")
                )
                document_metadata_vision = DocumentMetadata(
                    original_filename=safe_filename,
                    processed_at=datetime.now().isoformat(),
                    page_count=result_vision.get("page_count")
                )
                extracted_text_vision = result_vision.get("extracted_text", "")
                page_count_vision = result_vision.get("page_count", 0)
                pages_content_vision = []

                if extracted_text_vision and page_count_vision > 0:
                    # Try to split by the delimiter used in process_with_gemini_vision_ocr
                    page_texts_split = extracted_text_vision.split("\n\n") 
                    if len(page_texts_split) == page_count_vision:
                        for i, p_text in enumerate(page_texts_split):
                            pages_content_vision.append(PageContent(page_number=i + 1, text=p_text))
                    else:
                        logger.warning(f"Vision OCR: Page text split count ({len(page_texts_split)}) differs from page_count ({page_count_vision}). Creating a single page entry with all text.")
                        pages_content_vision.append(PageContent(page_number=1, text=extracted_text_vision))
                elif extracted_text_vision: # Has text but no page_count or page_count is 0
                    pages_content_vision.append(PageContent(page_number=1, text=extracted_text_vision))
                
                return OcrResponse(
                    document_metadata=document_metadata_vision,
                    processing_info=processing_info_vision,
                    pages=pages_content_vision,
                    full_text=extracted_text_vision,
                    tables=[], images=[], gemini_analysis=None
                )
        else: # Standard processing path (use_vision_ocr_bool is False)
            result_standard = process_pdf(
                file_path=temp_file_path,
                use_pypdf2=use_pypdf2_bool and PYPDF2_AVAILABLE,
                use_docling=use_docling_bool and DOCLING_AVAILABLE,
                use_gemini=use_gemini_bool and GEMINI_ENABLED,
                force_ocr=force_ocr_bool,
                language=language,
                gemini_model=gemini_model
            )

            if "error" in result_standard or "processing_info" not in result_standard : # Check for error or missing critical structure
                error_msg = result_standard.get("error", "Unknown error during standard processing or malformed result.")
                if "processing_info" not in result_standard:
                     error_msg += " (Result missing 'processing_info')"
                logger.error(f"Standard PDF processing failed or returned malformed result: {error_msg}")
                
                # Attempt to get processing time if available, otherwise None
                std_proc_time = result_standard.get("processing_info", {}).get("processing_time_seconds")

                error_processing_info = ProcessingInfo(
                    pypdf2_used=use_pypdf2_bool and PYPDF2_AVAILABLE,
                    docling_used=use_docling_bool and DOCLING_AVAILABLE,
                    gemini_used=use_gemini_bool and GEMINI_ENABLED,
                    force_ocr_used=force_ocr_bool,
                    status="error",
                    error_message=str(error_msg),
                    processing_time_seconds=std_proc_time
                )
                error_document_metadata = DocumentMetadata(
                    original_filename=safe_filename,
                    processed_at=datetime.now().isoformat(),
                    page_count=result_standard.get("document_metadata", {}).get("page_count")
                )
                return OcrResponse(
                    document_metadata=error_document_metadata,
                    processing_info=error_processing_info,
                    pages=[], full_text=result_standard.get("full_text", ""), tables=[], images=[], gemini_analysis=None
                )

            # Standard success path: result_standard is valid and from process_pdf
            processing_info = ProcessingInfo(
                pypdf2_used=result_standard["processing_info"]["pypdf2_used"],
                docling_used=result_standard["processing_info"]["docling_used"],
                gemini_used=result_standard["processing_info"]["gemini_used"],
                force_ocr_used=result_standard["processing_info"]["force_ocr_used"],
                processing_time_seconds=result_standard["processing_info"]["processing_time_seconds"],
                status=result_standard["processing_info"]["status"],
                error_message=result_standard["processing_info"].get("error_message")
            )
            document_metadata = DocumentMetadata(
                original_filename=result_standard["document_metadata"]["original_filename"],
                processed_at=result_standard["document_metadata"]["processed_at"],
                page_count=result_standard["document_metadata"].get("page_count"),
                source_format=result_standard["document_metadata"].get("source_format"),
                language=result_standard["document_metadata"].get("language"),
                title=result_standard["document_metadata"].get("title"),
                author=result_standard["document_metadata"].get("author"),
                creation_date=result_standard["document_metadata"].get("creation_date")
            )
            pages = []
            for page_data in result_standard.get("pages", []): # Use .get for safety
                page = PageContent(
                    page_number=page_data["page_number"],
                    text=page_data.get("text", ""), # Use .get for safety
                    tables=page_data.get("tables", []),
                    images=page_data.get("images", [])
                )
                pages.append(page)

            gemini_analysis = None
            if "gemini_analysis" in result_standard and result_standard["gemini_analysis"]:
                gemini_data = result_standard["gemini_analysis"]
                gemini_analysis = GeminiAnalysis(
                    summary=gemini_data.get("summary"),
                    key_points=gemini_data.get("key_points", []),
                    structured_data=gemini_data.get("structured_data"),
                    translation=gemini_data.get("translation"),
                    raw_response=gemini_data.get("raw_response")
                )
            tables = []
            for table_data in result_standard.get("tables", []): # Use .get for safety
                # Ensure raw_text is string, rows is list of lists
                # (existing robust conversion logic for tables in process_pdf and previous versions of ocr_upload can be reused or adapted here if needed)
                # For brevity, direct assignment, assuming process_pdf returns valid TableInfo-like dicts
                try:
                    # Ensure raw_text is string
                    if "raw_text" in table_data and not isinstance(table_data["raw_text"], str):
                        table_data["raw_text"] = json.dumps(table_data["raw_text"])
                    # Ensure rows is list of lists
                    if "rows" in table_data and (not isinstance(table_data["rows"], list) or (table_data["rows"] and not isinstance(table_data["rows"][0], list))):
                        # Basic fallback if rows format is not list of lists
                        logger.warning(f"Table rows for table '{table_data.get('table_id')}' not in expected list-of-lists format. Setting to empty list.")
                        table_data["rows"] = []

                    tables.append(TableInfo(**table_data))
                except Exception as e_table:
                    logger.warning(f"Skipping table due to parsing error: {e_table}. Table data: {table_data}")


            return OcrResponse(
                document_metadata=document_metadata,
                processing_info=processing_info,
                pages=pages,
                full_text=result_standard.get("full_text", ""), # Use .get for safety
                tables=tables,
                images=result_standard.get("images", []), # Use .get for safety
                gemini_analysis=gemini_analysis
            )

    except Exception as e:
        logger.exception(f"Error processing file {file.filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")
    finally:
        # 4. 清理临时文件
        if temp_file_path:
            cleanup_temp_file(temp_file_path)
            logger.info(f"Cleaned up temporary file: {temp_file_path}")

@app.get("/api/ocr/status")
async def ocr_status():
    """
    获取OCR服务的状态信息

    Returns:
        包含服务状态信息的JSON
    """
    return {
        "status": "running",
        "version": "1.3.0",  # 更新版本号，支持Excel、Word、文本文件
        "capabilities": {
            "pypdf2_available": PYPDF2_AVAILABLE,
            "docling_available": DOCLING_AVAILABLE,
            "gemini_enabled": GEMINI_ENABLED,
            "excel_processing_available": EXCEL_PROCESSING_AVAILABLE,
            "word_processing_available": WORD_PROCESSING_AVAILABLE,
            "default_gemini_model": DEFAULT_GEMINI_MODEL if GEMINI_ENABLED else None,
            "available_gemini_models": AVAILABLE_GEMINI_MODELS if GEMINI_ENABLED else [],
            "gemini_model_configs": GEMINI_MODELS if GEMINI_ENABLED else {},
            "supported_file_types": {
                "pdf": PYPDF2_AVAILABLE or DOCLING_AVAILABLE,
                "images": DOCLING_AVAILABLE,
                "excel": EXCEL_PROCESSING_AVAILABLE,
                "word": WORD_PROCESSING_AVAILABLE,
                "text": True,  # 文本文件总是支持
                "csv": True,  # CSV文件总是支持
                "markdown": True  # Markdown文件总是支持
            }
        },
        "cache_dir": os.path.abspath(LOCAL_CACHE_DIR),
        "upload_dir": os.path.abspath(UPLOAD_DIR)
    }

# --- Main execution (for local testing) ---
if __name__ == "__main__":
    import uvicorn
    # Make sure to install uvicorn and python-multipart:
    # pip install uvicorn python-multipart PyPDF2 google-generativeai Pillow
    # May also need to install docling if you want to use it:
    # pip install docling easyocr

    logger.info("Starting OCR microservice with Uvicorn on http://127.0.0.1:8012")
    uvicorn.run(app, host="127.0.0.1", port=8012)

# Example of how to run from command line:
# python ocr_service.py

# Example of how to test with curl:
# curl -X POST -F "file=@/path/to/your/document.pdf" http://127.0.0.1:8009/api/ocr/upload